import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to extract text from TXT file {file_path}: {e}")
            return ""
    except Exception as e:
        logger.error(f"Failed to extract text from TXT file {file_path}: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    try:
        import PyPDF2
        
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1} from {file_path}: {e}")
                    continue
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from PDF file {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX file {file_path}: {e}")
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower().strip('.')
    
    extractors = {
        'txt': extract_text_from_txt,
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'doc': extract_text_from_docx,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        logger.error(f"Unsupported file type: {file_type}")
        return ""
    
    if not Path(file_path).exists():
        logger.error(f"File not found: {file_path}")
        return ""
    
    return extractor(file_path)

# Made with Bob
